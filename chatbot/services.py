import time

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, ns
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches

data = {
    "data": "2023-03-08",
    "cliente": "Teste Cliente",
    "destinatario": "Público",
    "titulo": "Treinamento Avançado em Prestação de Serviços e Produtos",
    "objetivo": "Ampliar e aprimorar o conhecimento e habilidades técnicas dos profissionais envolvidos na prestação de serviços e produtos, visando otimizar o atendimento ao cliente e fortalecer a competitividade da organização.",
    "periodo": "10 dias corridos (com base na carga horária total)",
    "detalhamento_proposta": [
        {
            "titulo_etapa": "Módulo 1: Fundamentos de Atendimento ao Cliente",
            "carga_horaria_necessaria": "20 horas",
            "publico_alvo": "Profissionais de atendimento, gestores e supervisores",
            "objetivo_da_etapa": "Aprimorar as competências de comunicação, relacionamento interpessoal e resolução de conflitos, visando proporcionar um atendimento excepcional ao cliente.",
            "conteudo_programatico": [
                "Técnicas de comunicação eficaz",
                "Empatia e escuta ativa",
                "Gestão de conflitos e reclamações",
                "Atendimento personalizado e segmentado",
                "Avaliação e mensuração do atendimento",
            ],
        },
        {
            "titulo_etapa": "Módulo 1: Fundamentos de Atendimento ao Cliente",
            "carga_horaria_necessaria": "20 horas",
            "publico_alvo": "Profissionais de atendimento, gestores e supervisores",
            "objetivo_da_etapa": "Aprimorar as competências de comunicação, relacionamento interpessoal e resolução de conflitos, visando proporcionar um atendimento excepcional ao cliente.",
            "conteudo_programatico": [
                "Técnicas de comunicação eficaz",
                "Empatia e escuta ativa",
                "Gestão de conflitos e reclamações",
                "Atendimento personalizado e segmentado",
                "Avaliação e mensuração do atendimento",
            ],
        },
        {
            "titulo_etapa": "Módulo 1: Fundamentos de Atendimento ao Cliente",
            "carga_horaria_necessaria": "20 horas",
            "publico_alvo": "Profissionais de atendimento, gestores e supervisores",
            "objetivo_da_etapa": "Aprimorar as competências de comunicação, relacionamento interpessoal e resolução de conflitos, visando proporcionar um atendimento excepcional ao cliente.",
            "conteudo_programatico": [
                "Técnicas de comunicação eficaz",
                "Empatia e escuta ativa",
                "Gestão de conflitos e reclamações",
                "Atendimento personalizado e segmentado",
                "Avaliação e mensuração do atendimento",
            ],
        },
    ],
    "carga_horaria_total": "65 horas",
    "valor_investimento": "R$ 5.000,00",
}

REPLACEMENTS = {
    "{data}": "data",
    "{cliente}": "cliente",
    "{destinatario}": "destinatario",
    "{titulo}": "titulo",
    "{objetivo}": "objetivo",
    "{periodo}": "periodo",
    "{carga_horaria_total}": "carga_horaria_total",
    "{valor_investimento}": "valor_investimento",
}

SCHEMA_DETALHAMENTO_PROPOSTA = {
    "carga_horaria_necessaria": "Carga Horária",
    "publico_alvo": "Público-alvo",
    "objetivo_da_etapa": "Objetivo",
    "conteudo_programatico": "Conteúdo Programático (ementa)",
}


def replace_text_in_doc(doc, data_json):
    replace_paragraph(doc, data_json)
    if data_json.get("detalhamento_proposta"):
        insert_detalhamento_da_proposta(
            doc, data_json.get("detalhamento_proposta"))


def replace_paragraph(doc, data_json):
    for paragraph in doc.paragraphs:
        for key, value in REPLACEMENTS.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, data_json[value])


def insert_detalhamento_da_proposta(doc, detalhamento_proposta: list):
    table = doc.tables[2]
    for detalhe in detalhamento_proposta:
        for key, value in detalhe.items():
            if key == "titulo_etapa":
                insert_line_merge(table, value)
                continue
            insert_line_information(
                table, SCHEMA_DETALHAMENTO_PROPOSTA[key], value)

    first_row = table.rows[0]
    merged_cell = first_row.cells[0].merge(first_row.cells[1])
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    for row in table.rows:
        for cell in row.cells:
            set_cell_border(cell, top={"sz": 12, "val": "single", "color": "000000"},
                            bottom={"sz": 12, "val": "single",
                                    "color": "000000"},
                            left={"sz": 12, "val": "single", "color": "000000"},
                            right={"sz": 12, "val": "single", "color": "000000"})


def insert_line_merge(table, text):
    merged_row = table.add_row()
    merged_row.cells[0].text = text
    merged_cell = merged_row.cells[0].merge(merged_row.cells[1])
    merged_cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    merged_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    merged_cell.paragraphs[0].runs[0].bold = True


def insert_line_information(table, key, value):
    new_row = table.add_row()

    new_row.cells[0].text = key
    if isinstance(value, list):
        for item in value:
            new_row.cells[1].add_paragraph(f"• {item}")
    else:
        new_row.cells[1].text = value

    table_width = Inches(6)
    cell_1_width = int(table_width * 0.2)
    cell_2_width = int(table_width * 0.8)
    new_row.cells[0].width = cell_1_width
    new_row.cells[1].width = cell_2_width


def set_cell_border(cell, **kwargs):
    """
    Define bordas para uma célula. Os argumentos aceitos em kwargs são:
    top, left, bottom, right e os valores podem ser: 
    {"sz": tamanho, "val": "single", "color": "000000", "space": "0"}
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()

    # Criar o elemento <w:tcBorders> se não existir
    tcBorders = tcPr.find(qn('w:tcBorders'))
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    # Definir as bordas para cada lado (topo, esquerdo, inferior, direito)
    for edge in ('top', 'left', 'bottom', 'right'):
        if edge in kwargs:
            element = OxmlElement(f'w:{edge}')
            element.set(qn('w:val'), kwargs[edge].get("val", "single"))
            # Tamanho da borda (4 = 1/2 pt)
            element.set(qn('w:sz'), str(kwargs[edge].get("sz", 4)))
            element.set(qn('w:space'), str(kwargs[edge].get("space", 0)))
            element.set(qn('w:color'), kwargs[edge].get(
                "color", "000000"))  # Cor da borda
            tcBorders.append(element)


def generate_doc(data_json):
    doc = Document("modelo.docx")
    replace_text_in_doc(doc, data_json)
    file_name = f"{time.time()}.docx"
    doc.save(file_name)
    content_file = open(file_name, "rb")
    return content_file


if __name__ == '__main__':
    generate_doc(data)
